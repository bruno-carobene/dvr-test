# document_generator.py
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Inches 
import os
import sys

# Installa docxcompose se non presente (per Streamlit Cloud)
try:
    from docxcompose.composer import Composer
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "docxcompose==1.1.2", "--quiet"])
    from docxcompose.composer import Composer

# Database agenti chimici
db_chimico = {
    "Acidi per laboratori didattici": ["2-Medio", "H314"], 
    "Acido cloridrico": ["2-Medio", "H315 - H335"],
    "Acido solforico al 30%": ["2-Medio", "GHS05, H209, H314"], 
    "Acquaragia": ["2-Medio", "H304"],
    "Agenti pulitori sgrassanti": ["2-Medio", "H412, H304, H226, H336, H229"], 
    "Alcool etilico": ["1-Basso", "H225, H226, H319"],
    "Alghicida": ["3-Alto", "H314, H318"], 
    "Ammoniaca": ["3-Alto", "H315, H319"],
    "Antiadesivo siliconico": ["1-Basso", "H222, H229"], 
    "Anticorrosivo": ["2-Medio", "H22, H229,H315, H412"],
    "Antigelo": ["1-Basso", "H302, H304, H315, H318, H351"], 
    "Antigelo Permanente": ["1-Basso", "H302"],
    "Antiruggine": ["2-Medio", "H314, H315, H319"], 
    "Antiruggine liquido": ["1-Basso", "H226, H373, H315"],
    "Argon": ["1-Basso", "H280"], 
    "Azoto": ["1-Basso", "H281"],
    "Benzina": ["1-Basso", "H224, H304, H340, H350"], 
    "Blu di prussia": ["2-Medio", "H302, H315, H319"],
    "Candeggina": ["2-Medio", "H315, H319"], 
    "Catalizzatore vernici veicoli": ["2-Medio", "H226, H332, H304, H412"],
    "Collodio": ["1-Basso", "H226, H319, H335"], 
    "Correttore di pH": ["3-Alto", "H302, H314, H318, H400"],
    "Detergente disincrostante forni": ["2-Medio", "H315, H319"], 
    "Detergente igienizzante clima": ["1-Basso", "H225, H319"],
    "Detergente stoviglie a mano": ["3-Alto", "H302, H315, H318"], 
    "Detergente lavastoviglie": ["3-Alto", "H319, "],
    "Detergente lucidatura carrozzerie": ["1-Basso", "H225, H319, H336"], 
    "Detergente per pavimenti": ["1-Basso", "H315, H318"],
    "Detergente per superfici diluito": ["1-Basso", "-"], 
    "Detergente per WC": ["1-Basso", "H314, H335"],
    "Detergente speciale offset": ["1-Basso", "H226, H304, H336, H411"], 
    "Detersivo per lavatrice": ["1-Basso", "-"],
    "Diluente per inchiostri": ["2-Medio", "H226, H304, H335, H336"], 
    "Diluenti Nitro Antinebbia": ["3-Alto", "H225, H361d, H373"],
    "Flocculante": ["2-Medio", "H318"], 
    "Flussante": ["2-Medio", "H319, H336, H225"],
    "Fondo verniciatura veicoli": ["2-Medio", "H226, H314, H373, H412"], 
    "Fumi di saldatura": ["2-Medio", "n.a."],
    "Gasolio": ["1-Basso", "H226, H304, H332, H351, H411"], 
    "Glicole etilenico": ["1-Basso", "H302"],
    "Grasso lubrificante": ["1-Basso", "-"], 
    "Inchiostri per offset": ["2-Medio", "H315, H318, "],
    "Indurente vernici veicoli": ["2-Medio", "H226, H332, H317, H360, H412"], 
    "Legante basi verniciatura": ["2-Medio", "-"],
    "Loctite-401": ["1-Basso", "H315, H319, H335"], 
    "Lubrificanti spray (Svitol/Grasso)": ["1-Basso", "H223, H304, H411"],
    "Malta": ["2-Medio", "H318, H315, H317, H335"], 
    "Oli lubrificanti": ["1-Basso", "H315, H318, H336"],
    "Olio per impastare inchiostri": ["1-Basso", "H225, EUH066"], 
    "Pasta per riscontro": ["1-Basso", "-"],
    "Pasta per saldare i chip": ["2-Medio", "H302, H315, H317, H410"], 
    "Pittura ad acqua": ["2-Medio", "-"],
    "Polveri da molatura": ["2-Medio", "-"], 
    "Primer verniciatura veicoli": ["2-Medio", "H225, H315, H373, H412"],
    "Pulitore contatti elettrici": ["2-Medio", "H222, H315, H319, H411"], 
    "Reagenti": ["3-Alto", "Varia"],
    "Rivestimento trasparente veicoli": ["2-Medio", "H226, H317, H336, H412"], 
    "Sbloccante spray": ["1-Basso", "H223, H336, H229"],
    "Sepiolite": ["2-Medio", "H318, H302, H315"], 
    "Silicone spray": ["2-Medio", "H222, H315, H336, H411"],
    "Soluzione disinfettante": ["1-Basso", "H302, H318, H319, H336"], 
    "Solventi": ["3-Alto", "H304, "],
    "Tinta per capelli": ["2-Medio", "n.a."], 
    "Toner": ["1-Basso", "-"],
    "Total clean": ["1-Basso", "H319"], 
    "Vernice acqua spruzzo": ["2-Medio", "H319"],
    "Vernice spray": ["2-Medio", "H222, H319, H336, H411"], 
    "Vernici per offset": ["2-Medio", "H301, H314, H317, H411"]
}

def imposta_colore_cella(cella, colore_hex):
    """Imposta colore di sfondo cella"""
    xml_string = f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{colore_hex}"/>'
    shading_elm = parse_xml(xml_string)
    cella._tc.get_or_add_tcPr().append(shading_elm)

def sostituisci_mantieni_formato(paragrafo, placeholder, valore):
    """Sostituisce il placeholder mantenendo la formattazione"""
    if placeholder not in paragrafo.text:
        return False
    
    for run in paragrafo.runs:
        if placeholder in run.text:
            font_name = run.font.name
            font_size = run.font.size
            bold = run.font.bold
            italic = run.font.italic
            color = run.font.color.rgb if run.font.color else None
            
            run.text = run.text.replace(placeholder, str(valore))
            
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
            if bold:
                run.font.bold = bold
            if italic:
                run.font.italic = italic
            if color:
                run.font.color.rgb = color
            
            return True
    
    full_text = ''.join([r.text for r in paragrafo.runs])
    if placeholder in full_text:
        first_run = paragrafo.runs[0]
        for run in paragrafo.runs[1:]:
            run.text = ""
        first_run.text = full_text.replace(placeholder, str(valore))
        return True
    
    return False

def compila_segnaposto(doc, dati):
    """Sostituisce segnaposto mantenendo la formattazione"""
    for p in doc.paragraphs:
        for key, value in dati.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in p.text:
                sostituisci_mantieni_formato(p, placeholder, str(value))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in dati.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        for p in cell.paragraphs:
                            sostituisci_mantieni_formato(p, placeholder, str(value))

def inserisci_tabella_chimica(doc, segnaposto, lista_scelti, db):
    """Inserisce tabella agenti chimici"""
    for p in doc.paragraphs:
        if segnaposto in p.text:
            p.text = p.text.replace(segnaposto, "")
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            for i, t in enumerate(['Prodotto', 'Rischio', 'Classificazione']):
                hdr_cells[i].text = t
                imposta_colore_cella(hdr_cells[i], "D9D9D9")
                hdr_cells[i].paragraphs[0].runs[0].bold = True

            for nome_prod in lista_scelti:
                if nome_prod in db:
                    row_cells = table.add_row().cells
                    row_cells[0].text = nome_prod
                    rischio = db[nome_prod][0]
                    row_cells[1].text = rischio
                    row_cells[2].text = db[nome_prod][1]

                    if "Alto" in rischio: 
                        colore = "FF9999"
                    elif "Medio" in rischio: 
                        colore = "FFFFCC"
                    elif "Basso" in rischio: 
                        colore = "CCFFCC"
                    else: 
                        colore = "FFFFFF"

                    imposta_colore_cella(row_cells[1], colore)

            p._element.addnext(table._element)

def rimuovi_sommario_dinamico(doc):
    """Rimuove il campo sommario (TOC) dal documento"""
    for p in doc.paragraphs[:]:
        if p._element.xpath('.//w:fldChar'):
            p._element.getparent().remove(p._element)
            continue
        
        if any(testo in p.text.upper() for testo in [
            'NESSUNA VOCE DI SOMMARIO TROVATA',
            'AGGIORNA SOMMARIO',
            'SOMMARIO'
        ]):
            p._element.getparent().remove(p._element)

def aggiungi_sommario_statico(doc):
    """Aggiunge il sommario statico dopo la prima pagina"""
    # Trova l'interruzione pagina
    insert_point = None
    for i, p in enumerate(doc.paragraphs):
        if p._element.xpath('.//w:br[@w:type="page"]'):
            insert_point = p
            break
    
    if insert_point is None:
        # Se non c'è, aggiungiamo noi l'interruzione dopo il primo contenuto
        # Troviamo un punto appropriato (dopo la tabella)
        for p in reversed(doc.paragraphs):
            if p.text.strip():
                insert_point = p
                break
        if insert_point:
            insert_point.add_run().add_break()
    
    if insert_point is None:
        return
    
    # Inserisci titolo SOMMARIO
    p_titolo = doc.add_paragraph()
    insert_point._element.addnext(p_titolo._element)
    run = p_titolo.add_run("SOMMARIO")
    run.bold = True
    run.font.size = Pt(14)
    
    # Righe sommario
    sommario_items = [
        ("Introduzione", 0),
        ("Obiettivi del documento", 1),
        ("Chi ha partecipato alla redazione del documento", 1),
        ("Procedura di identificazione e analisi dei rischi e definizione dei controlli", 0),
        ("Identificazione dei centri/fonti di pericolo per la sicurezza e la salute dei lavoratori", 1),
        ("Identificazione dei lavoratori (o di terzi) esposti a rischi potenziali", 1),
        ("Valutazione dei rischi, dal punto di vista qualitativo e quantitativo", 1),
        ("Studio sulla possibilità di eliminare i rischi", 1),
        ("Programma delle misure ritenute opportune per garantire il miglioramento nel tempo dei livelli di sicurezza e procedura per l'attuazione", 1),
        ("Elenco dei pericoli considerati", 0),
        ("Criteri di quantificazione del rischio", 0),
        ("Prescrizioni legali", 0),
        ("Gestione del documento", 0),
        ("", 0),
        ("L'azienda", 0),
        ("Anagrafica aziendale", 1),
        ("Il Sistema di sicurezza aziendale", 1),
        ("Descrizione strutturale della sede di lavoro", 1),
        ("Descrizione generale dei locali", 2),
        ("Attività affidate a terzi", 2),
        ("Attività svolte presso terzi", 2),
        ("Attrezzature e agenti chimici impiegati", 0),
        ("", 0),
        ("Allegati", 0),
        ("Ambienti di lavoro", 1),
        ("Attrezzature", 1),
        ("Mansioni", 1),
    ]
    
    ultimo_para = p_titolo
    for testo, livello in sommario_items:
        p = doc.add_paragraph()
        ultimo_para._element.addnext(p._element)
        
        if testo:
            run = p.add_run(testo)
            if livello == 0:
                run.bold = True
                run.font.size = Pt(11)
            elif livello == 1:
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Pt(18)
            else:
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Pt(36)
        
        ultimo_para = p

def formatta_elenco_paragrafi(lista):
    """Formatta elenco come lista di stringhe"""
    if not lista:
        return []
    return [v.replace("_", " ").capitalize() for v in lista]

def inserisci_elenco_puntato(doc, segnaposto, lista_voci):
    """Sostituisce il segnaposto con un elenco puntato"""
    if not lista_voci:
        return
    
    for p in doc.paragraphs:
        if segnaposto in p.text:
            parent = p._element.getparent()
            idx = list(parent).index(p._element)
            
            p.text = p.text.replace(segnaposto, lista_voci[0] if lista_voci else "")
            
            try:
                p.style = 'List Bullet'
            except:
                pass
            
            for i, voce in enumerate(lista_voci[1:], start=1):
                new_p = doc.add_paragraph(voce)
                try:
                    new_p.style = 'List Bullet'
                except:
                    pass
                parent.insert(idx + i, new_p._element)
            
            break

def genera_dvr(azienda_data, ambienti, attrezzature, mansioni, agenti_chimici, templates_dir, logo_file=None, foto_ambienti=None):
    """
    Funzione principale che genera il documento DVR usando docxcompose
    """
    # Prepara dati
    data_di_oggi = datetime.now().strftime("%d/%m/%Y")
    azienda_data["DATA"] = data_di_oggi
    
    lista_ambienti = formatta_elenco_paragrafi(ambienti)
    lista_mansioni = formatta_elenco_paragrafi(mansioni)
    lista_attrezzature = formatta_elenco_paragrafi(attrezzature)
    lista_chimici = formatta_elenco_paragrafi(agenti_chimici)
    
    azienda_data["LISTA_AMBIENTI"] = "\n".join(lista_ambienti) if lista_ambienti else ""
    azienda_data["LISTA_MANSIONI"] = "\n".join(lista_mansioni) if lista_mansioni else ""
    azienda_data["LISTA_ATTREZZATURE"] = "\n".join(lista_attrezzature) if lista_attrezzature else ""
    azienda_data["LISTA_CHIMICI"] = "\n".join(lista_chimici) if lista_chimici else ""
    
    #template_path = os.path.join(templates_dir, 'Template_Base.docx')
   # --- Tutto questo blocco deve avere 4 spazi all'inizio di ogni riga ---
    template_path = os.path.join('templates', 'Template_Base.docx')
    
    # Controllo di debug
    if not os.path.exists(template_path):
        print(f"ERRORE: Non trovo il file in {os.path.abspath(template_path)}")
        # Se la cartella esiste, elenca i file, altrimenti segnala l'assenza della cartella
        if os.path.exists('templates'):
            print(f"File dentro templates: {os.listdir('templates')}")
        else:
            print("La cartella templates non esiste proprio nel container!")
    
    # 0. Carica template (DEVE essere allineato con la 'if' sopra)
    master_doc = Document(template_path)

    # 1. GESTIONE LOGO (Se caricato, lo inseriamo in cima)
    if logo_file:
        # Nota: Inseriamo l'immagine nel primo paragrafo (che dovrebbe essere vuoto o sopra il nome)
        # Se nel template hai rimosso l'immagine originale, lo mettiamo all'inizio
        p = master_doc.paragraphs[0]
        r = p.add_run()
        r.add_picture(logo_file, width=Inches(2.5)) # Ridimensiona a circa 6.3cm
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. Rimuovi sommario dinamico
    rimuovi_sommario_dinamico(master_doc)
    
    # 3. Compila segnaposti
    compila_segnaposto(master_doc, azienda_data)
    
    # 4. Inserisci elenchi
    inserisci_elenco_puntato(master_doc, "{{LISTA_ATTREZZATURE}}", lista_attrezzature)
    inserisci_elenco_puntato(master_doc, "{{LISTA_CHIMICI}}", lista_chimici)
    
    # 5. Inserisci tabella chimica
    inserisci_tabella_chimica(master_doc, "{{TABELLA_CHIMICA}}", agenti_chimici, db_chimico)

    # 6. GESTIONE FOTO AMBIENTI (Nuova pagina prima degli Allegati)
    if foto_ambienti:
        master_doc.add_page_break()
        titolo_foto = master_doc.add_heading("Ambienti di lavoro aziendali", level=1)
        titolo_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for foto in foto_ambienti:
            # Aggiunge l'immagine
            master_doc.add_picture(foto["file"], width=Inches(5))
            # Aggiunge la didascalia
            p_cap = master_doc.add_paragraph(foto["caption"])
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            master_doc.add_paragraph() # Spazio vuoto
    
    # 7. Aggiungi sommario statico
    aggiungi_sommario_statico(master_doc)
    
    # 8. Assembla moduli con Composer (SENZA salto pagina extra tra moduli)
    print("Assemblaggio moduli in corso...")
    composer = Composer(master_doc)
    
    # Raccogli tutti i moduli
    moduli_da_aggiungere = []
    for ambiente in ambienti:
        mod_path = os.path.join(templates_dir, f"{ambiente}.docx")
        if os.path.exists(mod_path):
            moduli_da_aggiungere.append(("ambiente", ambiente, mod_path))
    
    for att in attrezzature:
        mod_path = os.path.join(templates_dir, f"{att}.docx")
        if os.path.exists(mod_path):
            moduli_da_aggiungere.append(("attrezzatura", att, mod_path))
    
    for mans in mansioni:
        mod_path = os.path.join(templates_dir, f"{mans}.docx")
        if os.path.exists(mod_path):
            moduli_da_aggiungere.append(("mansione", mans, mod_path))
    
    # Aggiungi moduli (Composer gestisce automaticamente la paginazione)
    for tipo, nome, mod_path in moduli_da_aggiungere:
        try:
            doc = Document(mod_path)
            composer.append(doc)  # NON aggiungiamo salto pagina manualmente
            print(f"  ✓ Aggiunto {tipo}: {nome}")
        except Exception as e:
            print(f"  ✗ Errore con {nome}: {e}")
    
    # 8. Salva
    from io import BytesIO
    buffer = BytesIO()
    composer.save(buffer)
    buffer.seek(0)
    
    return buffer





